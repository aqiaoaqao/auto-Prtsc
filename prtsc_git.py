import time
import win32gui,win32ui,win32con,win32api

import os
from selenium import webdriver
from selenium.common.exceptions import NoAlertPresentException
import traceback

from docx import Document
from docx.shared import Inches,Cm


def windows_capture(filename):
	hwnd = 0                    
	hwndDC = win32gui.GetWindowDC(hwnd)
	mfcDC = win32ui.CreateDCFromHandle(hwndDC)
	saveDC = mfcDC.CreateCompatibleDC()
	saveBitMap = win32ui.CreateBitmap()
	MoniterDev = win32api.EnumDisplayMonitors(None,None)
	w = MoniterDev[0][2][2]
	h = MoniterDev[0][2][3]
	print (w,h)
	saveBitMap.CreateCompatibleBitmap(mfcDC,w,h)
	saveDC.SelectObject(saveBitMap)
	saveDC.BitBlt((0,0),(w,h),mfcDC,(0,0),win32con.SRCCOPY)
	saveBitMap.SaveBitmapFile(saveDC,filename)

def login(url,username_element,password_element,username,password,keybd):
	driver.get(url)

	driver.find_element_by_id(username_element).click()
	driver.find_element_by_id(username_element).clear()
	driver.find_element_by_id(username_element).send_keys(username)

	driver.find_element_by_id(password_element).click()
	driver.find_element_by_id(password_element).clear()
	driver.find_element_by_id(password_element).send_keys(password)

	driver.find_element_by_xpath(keybd).click()
	time.sleep(5)


def docx(ranks):
	document = Document('url_docx')
	run = document.tables[0].cell(ranks,0).paragraphs[0].add_run()
	run.add_picture('url_jpg',width = Cm(15))
	document.save('url_docx')

chromedriver = "driver_path"
os.environ["webdriver.ie.driver"] = chromedriver
driver = webdriver.Chrome()
driver.maximize_window()


try:
	alert1 = driver.switch_to.alert
except NoAlertPresentException as e:
	print("no alert")
	traceback.print_exc()
else:
	at_text1 = alert1.at_text
	print ("at_text:"+at_text1)
time.sleep(1)

login('  ')#此处传入的参数为：网址,用户名输入框元素名,密码输入框元素名,用户名,密码,登录框路径
time.sleep(2)
windows_capture("hah.jpg")
docx(0)
